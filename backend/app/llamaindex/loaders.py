import os
from typing import List
from llama_index.core import Document
from pypdf import PdfReader
from docx import Document as DocxReader

class DocumentLoader:
    @staticmethod
    def load_document(file_path: str, filename: str, file_type: str, document_id: str) -> List[Document]:
        """
        Loads PDF, DOCX, or TXT file into LlamaIndex Document instances.
        Attaches page metadata where available.
        """
        ext = os.path.splitext(filename)[1].lower()
        documents = []

        if ext == ".pdf" or "pdf" in file_type:
            reader = PdfReader(file_path)
            for idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    doc = Document(
                        text=text,
                        metadata={
                            "document_id": document_id,
                            "file_name": filename,
                            "page_label": str(idx + 1),
                            "page_number": idx + 1,
                            "file_type": file_type,
                        }
                    )
                    documents.append(doc)
            if not documents:
                # Fallback empty or unextractable PDF text
                documents.append(Document(text="[Empty or scanned PDF file]", metadata={"document_id": document_id, "file_name": filename}))

        elif ext == ".docx" or "word" in file_type or "docx" in file_type:
            docx_doc = DocxReader(file_path)
            full_text = []
            for p in docx_doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text)
            text_content = "\n\n".join(full_text) or "[Empty DOCX file]"
            documents.append(
                Document(
                    text=text_content,
                    metadata={
                        "document_id": document_id,
                        "file_name": filename,
                        "file_type": file_type,
                    }
                )
            )

        else: # Default TXT / Plain Text
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text_content = f.read()
            documents.append(
                Document(
                    text=text_content or "[Empty text file]",
                    metadata={
                        "document_id": document_id,
                        "file_name": filename,
                        "file_type": file_type,
                    }
                )
            )

        return documents
